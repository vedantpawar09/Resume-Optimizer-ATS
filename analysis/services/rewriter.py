"""
Builds the final "ATS Optimized Resume" output files.

Layout-preservation strategy:
 - If the original file was a DOCX: we re-open the ORIGINAL docx and only
   replace the *text* inside existing paragraphs/runs, section by section,
   using the AI's rewritten_sections dict. Fonts, spacing, headings order,
   bullet styles, margins, and page layout all come from the untouched
   original document - we never delete/re-create paragraphs, only swap
   their text content. This is a best-effort alignment (by section, then
   by line count) and works very well for typical single-column resumes.
 - If the original file was a PDF: PDFs do not have an editable "paragraph"
   model, so pixel-perfect in-place text replacement is not reliable. We
   generate a clean, professionally formatted PDF/DOCX using the SAME
   section order and heading names as the original, via ReportLab/DOCX,
   which preserves structure (not literal pixel layout).
 - A plain TXT export is always produced from the rewritten sections.
"""
import io
import logging
import re

import docx
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

logger = logging.getLogger('resume_optimizer')


def _normalize_key(key: str) -> str:
    return key.strip().lower().replace(' ', '_').replace(':', '')


_SKILLS_HINTS = ('skill', 'tool', 'competenc')
_MULTI_ENTRY_HINTS = ('experience', 'employment', 'work', 'education', 'project', 'certif', 'licen')


def _is_skills_key(key: str) -> bool:
    return any(hint in key.lower() for hint in _SKILLS_HINTS)


def _is_multi_entry_key(key: str) -> bool:
    return any(hint in key.lower() for hint in _MULTI_ENTRY_HINTS)


def _looks_like_meta_line(line: str) -> bool:
    """A 'Company | Start - End' style line: has a separator or a year."""
    return '|' in line or bool(re.search(r'\b(19|20)\d{2}\b', line))


def parse_entries(text: str) -> list:
    """Split a section's text into individual position/degree entries
    (separated by blank lines), each broken into an optional grey meta line
    ('Organization | Dates'), a title line, and remaining body lines. Falls
    back gracefully for sections that aren't actually multi-entry."""
    blocks = [b.strip('\n') for b in re.split(r'\n\s*\n', text.strip()) if b.strip()]
    entries = []
    for block in blocks:
        lines = [ln.strip() for ln in block.split('\n') if ln.strip()]
        if not lines:
            continue
        if len(lines) >= 2 and _looks_like_meta_line(lines[0]):
            entries.append({'meta': lines[0], 'title': lines[1], 'body': lines[2:]})
        else:
            entries.append({'meta': '', 'title': lines[0], 'body': lines[1:]})
    return entries


def extract_contact_line(resume) -> str:
    """Pull an email / phone / LinkedIn contact line out of the resume's
    header block (falling back to the first few hundred characters of raw
    text), for display under the candidate's name in the generated resume.

    Some PDF templates render the header in a visual small-caps/all-caps
    style, but store the ACTUAL characters as uppercase in the text layer
    (e.g. "VEDANTPAWAR0915@GMAIL.COM"). Since emails and URLs are
    conventionally lowercase regardless of how they were styled in the
    original document, we normalize case on extraction rather than
    reproducing whatever capitalization the source PDF happened to have."""
    header = ''
    if getattr(resume, 'structured_sections', None):
        header = resume.structured_sections.get('header', '') or ''
    text = header if header.strip() else (resume.raw_text or '')[:600]

    parts = []
    email_m = re.search(r'[\w.\-]+@[\w.\-]+\.\w+', text)
    if email_m:
        parts.append(email_m.group(0).lower())
    phone_m = re.search(r'\+?\d[\d\-\s().]{7,}\d', text)
    if phone_m:
        parts.append(re.sub(r'\s+', ' ', phone_m.group(0)).strip())
    linkedin_m = re.search(r'linkedin\.com/\S+', text, re.I)
    if linkedin_m:
        parts.append(linkedin_m.group(0).rstrip('.,)').lower())
    return '   |   '.join(parts)


def rewrite_docx_in_place(original_file_path: str, structured_sections: dict,
                           rewritten_sections: dict, output_path: str) -> str:
    """Replace section text inside the original DOCX, preserving all styling."""
    document = docx.Document(original_file_path)
    paragraphs = document.paragraphs

    # Build section -> paragraph index ranges by re-scanning headings, mirroring
    # the same heading detection logic used when the resume was first parsed.
    section_ranges = []
    current_section = 'header'
    start_idx = 0
    for i, p in enumerate(paragraphs):
        norm = _normalize_key(p.text)
        if norm and norm in {_normalize_key(k) for k in structured_sections.keys()} and len(p.text.strip()) < 40:
            section_ranges.append((current_section, start_idx, i))
            current_section = norm
            start_idx = i + 1
    section_ranges.append((current_section, start_idx, len(paragraphs)))

    rewritten_lookup = {_normalize_key(k): v for k, v in rewritten_sections.items()}

    for section_name, start, end in section_ranges:
        new_text = rewritten_lookup.get(section_name)
        if not new_text:
            continue
        new_lines = [ln for ln in new_text.split('\n') if ln.strip()]
        body_paragraphs = [p for p in paragraphs[start:end] if p.text.strip()]

        if not body_paragraphs:
            continue

        # Distribute new lines across the existing paragraphs, preserving each
        # paragraph's original style/run-formatting; extra lines are appended
        # to the last paragraph, missing lines leave trailing paragraphs blank.
        for idx, para in enumerate(body_paragraphs):
            if idx < len(new_lines):
                text = new_lines[idx]
            elif idx == len(body_paragraphs) - 1 and len(new_lines) > len(body_paragraphs):
                text = ' '.join(new_lines[idx:])
            else:
                text = ''
            _set_paragraph_text_preserve_style(para, text)

        # if there are more new lines than paragraphs, append them using the
        # last paragraph's style as a template
        if len(new_lines) > len(body_paragraphs) and body_paragraphs:
            template = body_paragraphs[-1]
            for extra in new_lines[len(body_paragraphs):]:
                new_para = document.add_paragraph(style=template.style)
                if template.runs:
                    run = new_para.add_run(extra)
                    run.font.size = template.runs[0].font.size
                    run.font.name = template.runs[0].font.name
                    run.bold = template.runs[0].bold
                else:
                    new_para.add_run(extra)

    document.save(output_path)
    return output_path


def _set_paragraph_text_preserve_style(paragraph, new_text: str):
    """Replace a paragraph's visible text while keeping its first run's
    formatting (font, size, bold, color, etc.)."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first_run = paragraph.runs[0]
    first_run.text = new_text
    for extra_run in paragraph.runs[1:]:
        extra_run.text = ''


def build_clean_docx(rewritten_sections: dict, candidate_name: str, output_path: str,
                      role_title: str = '', contact_line: str = '') -> str:
    """Generate a clean, single-column DOCX from scratch (the app's default
    template): a centered name/role/contact header, then bold-caps section
    headings with a thin rule after each section's content, and structured
    entries (grey 'Organization | Dates' line, bold title, description) for
    experience/education-style sections."""
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    candidate_name = (candidate_name or 'Resume').strip()[:45]
    role_title = (role_title or '').strip()[:60]
    contact_line = (contact_line or '').strip()[:140]

    INK = RGBColor(0x1F, 0x29, 0x37)
    META_GREY = RGBColor(0x6B, 0x72, 0x80)
    PILL_BG = 'F3F4F6'

    document = docx.Document()
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = INK

    for sec in document.sections:
        sec.top_margin = Inches(0.65)
        sec.bottom_margin = Inches(0.65)
        sec.left_margin = Inches(0.8)
        sec.right_margin = Inches(0.8)

    def add_divider(paragraph, size='10', color='9CA3AF'):
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), size)
        bottom.set(qn('w:space'), '6')
        bottom.set(qn('w:color'), color)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def shade_cell(cell, hex_color):
        cell_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        cell_pr.append(shd)

    def set_cell_borders_none(cell):
        cell_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'nil')
            borders.append(el)
        cell_pr.append(borders)

    # ---- Centered header ----
    name_p = document.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(candidate_name.upper())
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.name = 'Calibri'
    name_run.font.color.rgb = INK
    name_p.paragraph_format.space_after = Pt(2)

    if role_title:
        role_p = document.add_paragraph()
        role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role_run = role_p.add_run(role_title)
        role_run.font.size = Pt(14)
        role_run.font.name = 'Calibri'
        role_run.font.color.rgb = INK
        role_p.paragraph_format.space_after = Pt(8)

    if contact_line:
        contact_p = document.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_p.add_run(contact_line)
        contact_run.font.size = Pt(10)
        contact_run.font.name = 'Calibri'
        contact_run.font.color.rgb = META_GREY
        contact_p.paragraph_format.space_after = Pt(8)
        add_divider(contact_p, size='14', color='1F2937')

    document.add_paragraph().paragraph_format.space_after = Pt(2)

    # ---- Sections ----
    section_items = [(k, t) for k, t in rewritten_sections.items() if t and t.strip()]
    for idx, (key, text) in enumerate(section_items):
        heading_p = document.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(10)
        heading_p.paragraph_format.space_after = Pt(8)
        heading_run = heading_p.add_run(key.replace('_', ' ').upper())
        heading_run.font.size = Pt(13)
        heading_run.font.bold = True
        heading_run.font.name = 'Calibri'
        heading_run.font.color.rgb = INK

        if _is_skills_key(key):
            tokens = [t.strip() for t in re.split(r'[,\n;•]', text) if t.strip()]
            n_cols = 3 if len(tokens) > 4 else max(1, len(tokens))
            rows_needed = (len(tokens) + n_cols - 1) // n_cols
            skills_table = document.add_table(rows=rows_needed, cols=n_cols)
            for i, tok in enumerate(tokens):
                r, c = divmod(i, n_cols)
                cell = skills_table.rows[r].cells[c]
                set_cell_borders_none(cell)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(f'•  {tok}')
                run.font.size = Pt(10.3)
                run.font.name = 'Calibri'
                run.font.color.rgb = INK
            document.add_paragraph().paragraph_format.space_after = Pt(4)

        elif _is_multi_entry_key(key):
            for entry in parse_entries(text):
                if entry['meta']:
                    meta_p = document.add_paragraph()
                    meta_p.paragraph_format.space_after = Pt(1)
                    meta_run = meta_p.add_run(entry['meta'])
                    meta_run.font.size = Pt(10)
                    meta_run.font.name = 'Calibri'
                    meta_run.font.color.rgb = META_GREY
                if entry['title']:
                    title_p = document.add_paragraph()
                    title_p.paragraph_format.space_after = Pt(3)
                    title_run = title_p.add_run(entry['title'])
                    title_run.font.size = Pt(11)
                    title_run.font.bold = True
                    title_run.font.name = 'Calibri'
                    title_run.font.color.rgb = INK
                for line in entry['body']:
                    is_bullet = line.startswith(('-', '•'))
                    body_line = line.lstrip('-• ').strip()
                    body_p = document.add_paragraph(style='List Bullet' if is_bullet else None)
                    body_p.paragraph_format.space_after = Pt(4)
                    body_p.paragraph_format.line_spacing = 1.2
                    body_run = body_p.add_run(body_line)
                    body_run.font.size = Pt(10.3)
                    body_run.font.name = 'Calibri'
                    body_run.font.color.rgb = INK
                document.add_paragraph().paragraph_format.space_after = Pt(6)

        else:
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                is_bullet = line.startswith(('-', '•'))
                body_line = line.lstrip('-• ').strip()
                body_p = document.add_paragraph(style='List Bullet' if is_bullet else None)
                body_p.paragraph_format.space_after = Pt(4)
                body_p.paragraph_format.line_spacing = 1.25
                body_run = body_p.add_run(body_line)
                body_run.font.size = Pt(10.5)
                body_run.font.name = 'Calibri'
                body_run.font.color.rgb = INK

        if idx < len(section_items) - 1:
            divider_p = document.add_paragraph()
            divider_p.paragraph_format.space_after = Pt(10)
            add_divider(divider_p, size='8', color='D1D5DB')

    document.save(output_path)
    return output_path


def build_clean_pdf(rewritten_sections: dict, candidate_name: str, output_path: str,
                     role_title: str = '', contact_line: str = '') -> str:
    """Generate a clean, single-column PDF resume from scratch (the app's
    default template): a centered bold name, role subtitle, and contact
    line, a thin rule, then bold-caps section headings each followed by
    their content and a thin divider - with structured entries (grey
    'Organization | Dates' line, bold title, bullets) for experience and
    education-style sections, and a 3-column bullet grid for skills.

    Uses a plain single Frame (via SimpleDocTemplate) so pagination is
    entirely automatic and safe for resumes of any length - no custom
    multi-frame bookkeeping required."""
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import HRFlowable, Table, TableStyle

    candidate_name = (candidate_name or 'Resume').strip()[:45]
    role_title = (role_title or '').strip()[:60]
    contact_line = (contact_line or '').strip()[:140]

    INK = HexColor('#1F2937')
    META_GREY = HexColor('#6B7280')
    RULE_DARK = HexColor('#1F2937')
    RULE_LIGHT = HexColor('#D1D5DB')
    FOOTER_BAR = HexColor('#4B5563')

    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        'NameStyle', parent=styles['Title'], alignment=TA_CENTER,
        fontName='Helvetica-Bold', fontSize=25, textColor=INK, leading=28, spaceAfter=4,
    )
    role_style = ParagraphStyle(
        'RoleStyle', parent=styles['Normal'], alignment=TA_CENTER,
        fontName='Helvetica', fontSize=14, textColor=INK, spaceAfter=8,
    )
    contact_style = ParagraphStyle(
        'ContactStyle', parent=styles['Normal'], alignment=TA_CENTER,
        fontName='Helvetica', fontSize=10, textColor=META_GREY,
    )
    heading_style = ParagraphStyle(
        'SectionHeading', parent=styles['Heading2'], fontName='Helvetica-Bold',
        fontSize=13.5, textColor=INK, spaceBefore=0, spaceAfter=8, leading=16,
        letterSpacing=0.8,
    )
    meta_style = ParagraphStyle(
        'Meta', parent=styles['Normal'], fontName='Helvetica', fontSize=10,
        textColor=META_GREY, spaceAfter=1,
    )
    title_style = ParagraphStyle(
        'EntryTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11,
        textColor=INK, spaceAfter=4,
    )
    body_style = ParagraphStyle(
        'Body', parent=styles['Normal'], fontName='Helvetica', fontSize=10.3,
        textColor=INK, leading=15.5, spaceAfter=8,
    )
    bullet_style = ParagraphStyle(
        'Bullet', parent=body_style, leftIndent=14, bulletIndent=2, spaceAfter=5,
    )
    skill_style = ParagraphStyle(
        'Skill', parent=styles['Normal'], fontName='Helvetica', fontSize=10.3,
        textColor=INK, spaceAfter=6,
    )

    doc = SimpleDocTemplate(output_path, pagesize=LETTER,
                             leftMargin=0.8 * inch, rightMargin=0.8 * inch,
                             topMargin=0.65 * inch, bottomMargin=0.55 * inch)

    def footer(canvas, _doc):
        canvas.saveState()
        canvas.setFillColor(FOOTER_BAR)
        canvas.rect(0, 0, LETTER[0], 0.22 * inch, fill=1, stroke=0)
        canvas.restoreState()

    story = [Paragraph(candidate_name.upper(), name_style)]
    if role_title:
        story.append(Paragraph(role_title, role_style))
    if contact_line:
        story.append(Paragraph(contact_line, contact_style))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width='100%', thickness=1.4, color=RULE_DARK, spaceAfter=14))

    section_items = [(k, t) for k, t in rewritten_sections.items() if t and t.strip()]
    for idx, (key, text) in enumerate(section_items):
        story.append(Paragraph(key.replace('_', ' ').upper(), heading_style))

        if _is_skills_key(key):
            tokens = [t.strip() for t in re.split(r'[,\n;•]', text) if t.strip()]
            n_cols = 3 if len(tokens) > 4 else max(1, len(tokens))
            rows = [tokens[i:i + n_cols] for i in range(0, len(tokens), n_cols)]
            col_w = doc.width / n_cols
            table_data = [[Paragraph(f'•  {tok}', skill_style) for tok in row] for row in rows]
            skills_table = Table(table_data, colWidths=[col_w] * n_cols)
            skills_table.setStyle(TableStyle([
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 1),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(skills_table)

        elif _is_multi_entry_key(key):
            for entry in parse_entries(text):
                if entry['meta']:
                    story.append(Paragraph(entry['meta'], meta_style))
                if entry['title']:
                    story.append(Paragraph(entry['title'], title_style))
                if entry['body']:
                    items = [ListItem(Paragraph(ln.lstrip('-• ').strip(), bullet_style))
                             for ln in entry['body']]
                    story.append(ListFlowable(items, bulletType='bullet', leftIndent=14, bulletFontSize=7))
                story.append(Spacer(1, 8))

        else:
            bullets = [ln.strip('- ').strip() for ln in text.split('\n') if ln.strip()]
            if len(bullets) > 1:
                items = [ListItem(Paragraph(b, bullet_style)) for b in bullets]
                story.append(ListFlowable(items, bulletType='bullet', leftIndent=14, bulletFontSize=7))
            elif bullets:
                story.append(Paragraph(bullets[0], body_style))

        if idx < len(section_items) - 1:
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width='100%', thickness=0.7, color=RULE_LIGHT, spaceAfter=12))
        else:
            story.append(Spacer(1, 10))

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return output_path


def assemble_full_resume_text(rewritten_sections: dict, candidate_name: str = '') -> str:
    """Join every rewritten section into one continuous, readable document —
    used for the 'Complete Updated Resume' preview on the results/compare pages."""
    parts = []
    if candidate_name:
        parts.append(candidate_name.strip())
        parts.append('')
    for section, text in rewritten_sections.items():
        if not text or not text.strip():
            continue
        parts.append(section.replace('_', ' ').upper())
        parts.append(text.strip())
        parts.append('')
    return '\n'.join(parts).strip()


def build_txt(rewritten_sections: dict, candidate_name: str, output_path: str) -> str:
    lines = [candidate_name or 'Resume', '=' * 40, '']
    for section, text in rewritten_sections.items():
        lines.append(section.replace('_', ' ').upper())
        lines.append('-' * 30)
        lines.append(text.strip())
        lines.append('')
    content = '\n'.join(lines)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    return output_path