"""
File parsing utilities for resumes and job descriptions.

Supports PDF (via pdfplumber, falling back to PyMuPDF for tricky files)
and DOCX (via python-docx). Produces both a flat plain-text version and a
best-effort structured section split, used later by the ATS analyzer and
by the DOCX/PDF rewriters that must preserve the original layout.
"""
import io
import logging
import re

import pdfplumber
import fitz  # PyMuPDF
import docx

logger = logging.getLogger('resume_optimizer')

SECTION_HEADERS = [
    'summary', 'objective', 'profile', 'experience', 'work experience',
    'professional experience', 'professional summary', 'career objective',
    'relevant experience', 'work history', 'employment history', 'projects',
    'personal projects', 'academic projects', 'education', 'skills',
    'technical skills', 'core competencies', 'key skills', 'certifications',
    'certificates', 'licenses', 'achievements', 'accomplishments', 'awards',
    'languages', 'publications', 'volunteer', 'volunteering', 'interests',
    'hobbies', 'references', 'contact',
]


class ParsingError(Exception):
    pass


def extract_text_from_pdf(file_obj) -> str:
    """Extract text from a PDF file object using pdfplumber, with a PyMuPDF fallback."""
    text_parts = []
    try:
        file_obj.seek(0)
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ''
                text_parts.append(page_text)
        text = '\n'.join(text_parts).strip()
        if text:
            return text
    except Exception as exc:  # noqa: BLE001
        logger.warning("pdfplumber failed (%s), falling back to PyMuPDF", exc)

    try:
        file_obj.seek(0)
        data = file_obj.read()
        doc = fitz.open(stream=data, filetype='pdf')
        text = '\n'.join(page.get_text() for page in doc)
        doc.close()
        return text.strip()
    except Exception as exc:  # noqa: BLE001
        raise ParsingError(f"Could not extract text from PDF: {exc}") from exc


def extract_text_from_docx(file_obj) -> str:
    """Extract text from a DOCX file object using python-docx, including tables."""
    try:
        file_obj.seek(0)
        document = docx.Document(file_obj)
        chunks = []
        for para in document.paragraphs:
            if para.text.strip():
                chunks.append(para.text)
        for table in document.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    chunks.append(row_text)
        return '\n'.join(chunks).strip()
    except Exception as exc:  # noqa: BLE001
        raise ParsingError(f"Could not extract text from DOCX: {exc}") from exc


def extract_text(file_obj, file_type: str) -> str:
    file_type = file_type.lower().lstrip('.')
    if file_type == 'pdf':
        return extract_text_from_pdf(file_obj)
    if file_type == 'docx':
        return extract_text_from_docx(file_obj)
    raise ParsingError(f"Unsupported file type: {file_type}")


def split_into_sections(text: str) -> dict:
    """
    Best-effort split of resume plain text into named sections based on
    common resume section headings. Falls back to putting everything
    under 'full_text' if no headings are detected.
    """
    lines = text.splitlines()
    sections = {}
    current_key = 'header'
    buffer = []

    def flush():
        if buffer:
            sections.setdefault(current_key, '')
            sections[current_key] = (sections[current_key] + '\n' + '\n'.join(buffer)).strip()

    for line in lines:
        stripped = line.strip()
        lower = stripped.lower().strip(':').strip()
        is_header = (
            0 < len(stripped) < 40
            and any(re.search(rf'\b{re.escape(h)}\b', lower) for h in SECTION_HEADERS)
        )
        if is_header:
            flush()
            buffer = []
            current_key = re.sub(r'[^a-z]+', '_', lower).strip('_') or 'section'
        else:
            buffer.append(line)
    flush()

    if len(sections) <= 1:
        sections = {'full_text': text}
    return sections
